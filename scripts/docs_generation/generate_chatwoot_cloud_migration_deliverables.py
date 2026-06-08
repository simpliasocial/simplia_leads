from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = REPO_ROOT / "docs"
MIGRATIONS_DIR = REPO_ROOT / "supabase" / "migrations"

SOURCE_TECH_DOCX = DOCS_DIR / "Implementacion_Supabase_SimpliaLeads.docx"
SOURCE_SQL = DOCS_DIR / "supabase_replicacion_simpliale_solo_public_cw.sql"

UPDATED_TECH_DOCX = DOCS_DIR / "Implementacion_Supabase_SimpliaLeads_ISO10013_v2_1.docx"
UPDATED_SQL = DOCS_DIR / "supabase_replicacion_simpliale_solo_public_cw_actualizado_20260528.sql"
ONBOARDING_DOCX = DOCS_DIR / "Onboarding_Migracion_Chatwoot_Cloud_ISO10013.docx"

DOC_DATE = date.today().strftime("%d/%m/%Y")

COLORS = {
    "blue": "274690",
    "navy": "0f2344",
    "slate": "64748b",
    "line": "d9e2ef",
    "light": "f8fafc",
    "green": "0a9b6f",
    "mint": "dff7ec",
    "yellow": "fff7ed",
    "red": "dc2626",
    "white": "ffffff",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(document: Document, title: str, footer_text: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Title", 20, COLORS["blue"]),
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(footer_text)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(21)
    run.font.color.rgb = rgb(COLORS["blue"])

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(COLORS["slate"])
    document.add_paragraph()


def add_table(document: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True

    for row_index, values in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            run.font.name = "Arial"
            run.font.size = Pt(7.6 if row_index else 8)
            if row_index == 0:
                run.bold = True
                run.font.color.rgb = rgb(COLORS["white"])
                shade(cell, COLORS["blue"])
            else:
                shade(cell, COLORS["white"] if row_index % 2 else COLORS["light"])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths[: len(row.cells)]):
                row.cells[index].width = Inches(width)
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_callout(document: Document, title: str, body: str, tone: str = "blue") -> None:
    fill = COLORS["mint"] if tone == "green" else COLORS["yellow"] if tone == "yellow" else "eef4ff"
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8.8)
    run.font.color.rgb = rgb(COLORS["navy"])
    paragraph.add_run("\n")
    run = paragraph.add_run(body)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["navy"])
    document.add_paragraph()


def add_code(document: Document, text: str, title: str | None = None) -> None:
    if title:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.4)
        run.font.color.rgb = rgb(COLORS["blue"])
    clean = dedent(text).strip("\n")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    for line in clean.splitlines():
        run = paragraph.add_run(line.rstrip() + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(6.9)
        run.font.color.rgb = rgb(COLORS["navy"])


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def migration_text(name: str) -> str:
    path = MIGRATIONS_DIR / name
    if not path.exists():
        raise FileNotFoundError(path)
    return read_text(path).strip()


def n8n_chat_histories_migration_sql() -> str:
    return dedent(
        r"""
        -- ========================================================================
        -- INICIO BLOQUE 10: migración opcional public.n8n_chat_histories hacia cw
        -- DONDE SE EJECUTA: Supabase > SQL Editor
        -- INSTRUCCION: ejecutar solo si el proyecto origen tiene public.n8n_chat_histories.
        -- IMPORTANTE: este bloque NO borra public.n8n_chat_histories; la deja como fuente auditada.
        -- ========================================================================

        do $$
        begin
            if to_regclass('public.n8n_chat_histories') is null then
                raise notice 'public.n8n_chat_histories no existe; se omite migración opcional.';
            end if;
        end $$;

        insert into cw.import_batches (
            id,
            file_name,
            file_type,
            source_system,
            row_count,
            valid_count,
            skipped_count,
            create_count,
            update_count,
            status,
            mapping,
            stats,
            started_at,
            finished_at
        )
        select
            '00000000-0000-0000-0000-000000000908'::uuid,
            'public.n8n_chat_histories',
            'sql',
            'n8n_chat_histories',
            count(*),
            count(*) filter (
                where btrim(coalesce(
                    message->>'content',
                    message->>'text',
                    message #>> '{kwargs,content}',
                    message #>> '{data,content}',
                    message::text,
                    ''
                )) <> ''
            ),
            count(*) filter (
                where btrim(coalesce(
                    message->>'content',
                    message->>'text',
                    message #>> '{kwargs,content}',
                    message #>> '{data,content}',
                    message::text,
                    ''
                )) = ''
            ),
            count(distinct session_id),
            0,
            case
                when count(*) filter (
                    where btrim(coalesce(
                        message->>'content',
                        message->>'text',
                        message #>> '{kwargs,content}',
                        message #>> '{data,content}',
                        message::text,
                        ''
                    )) = ''
                ) > 0 then 'partial'
                else 'success'
            end,
            jsonb_build_object(
                'source_table', 'public.n8n_chat_histories',
                'identity_order', array['celular', 'correo', 'nombre', 'session_id'],
                'conversation_grain', 'session_id'
            ),
            jsonb_build_object(
                'source_table', 'public.n8n_chat_histories',
                'migrated_at', now()
            ),
            now(),
            now()
        from public.n8n_chat_histories
        where to_regclass('public.n8n_chat_histories') is not null
        on conflict (id) do update set
            row_count = excluded.row_count,
            valid_count = excluded.valid_count,
            skipped_count = excluded.skipped_count,
            create_count = excluded.create_count,
            status = excluded.status,
            mapping = excluded.mapping,
            stats = excluded.stats,
            finished_at = now();

        with source_rows as (
            select
                h.*,
                nullif(regexp_replace(coalesce(h.celular, ''), '\D', '', 'g'), '') as phone_digits,
                btrim(coalesce(
                    h.message->>'content',
                    h.message->>'text',
                    h.message #>> '{kwargs,content}',
                    h.message #>> '{data,content}',
                    h.message::text,
                    ''
                )) as message_content,
                lower(btrim(coalesce(
                    h.message->>'type',
                    h.message->>'role',
                    h.message #>> '{kwargs,type}',
                    h.message #>> '{data,type}',
                    ''
                ))) as message_kind
            from public.n8n_chat_histories h
            where to_regclass('public.n8n_chat_histories') is not null
        ),
        enriched as (
            select
                *,
                coalesce(
                    'phone:' || nullif(phone_digits, ''),
                    'email:' || nullif(lower(btrim(coalesce(correo, ''))), ''),
                    'name:' || nullif(lower(regexp_replace(btrim(coalesce(nombre, '')), '\s+', '_', 'g')), ''),
                    'session:' || session_id
                ) as identity_key,
                -1000000000::bigint - min(id)::bigint over (partition by session_id) as conversation_id,
                -2000000000::bigint - min(id)::bigint over (
                    partition by coalesce(
                        'phone:' || nullif(phone_digits, ''),
                        'email:' || nullif(lower(btrim(coalesce(correo, ''))), ''),
                        'name:' || nullif(lower(regexp_replace(btrim(coalesce(nombre, '')), '\s+', '_', 'g')), ''),
                        'session:' || session_id
                    )
                ) as contact_id,
                -3000000000::bigint - id::bigint as message_id,
                nullif(trim(both '_' from regexp_replace(lower(coalesce(pipeline_etapa, '')), '[^a-z0-9]+', '_', 'g')), '') as stage_label
            from source_rows
        ),
        contacts as (
            select distinct on (identity_key)
                identity_key,
                contact_id,
                nombre,
                correo,
                celular,
                canal,
                perfil_url,
                created_at,
                id,
                message
            from enriched
            order by identity_key, created_at nulls last, id
        )
        insert into cw.contacts_current (
            chatwoot_contact_id,
            lead_identity_key,
            identifier,
            name,
            phone_number,
            email,
            additional_attributes,
            custom_attributes,
            created_at_chatwoot,
            last_activity_at_chatwoot,
            first_seen_at,
            last_seen_at,
            raw_payload,
            updated_at
        )
        select
            contact_id,
            identity_key,
            identity_key,
            coalesce(nullif(btrim(nombre), ''), 'Sin Nombre'),
            nullif(btrim(celular), ''),
            nullif(btrim(correo), ''),
            jsonb_build_object('imported', true, 'source_system', 'n8n_chat_histories'),
            jsonb_strip_nulls(jsonb_build_object(
                'nombre_completo', nullif(btrim(nombre), ''),
                'celular', nullif(btrim(celular), ''),
                'correo', nullif(btrim(correo), ''),
                'canal', nullif(btrim(canal), ''),
                'perfil_url', nullif(btrim(perfil_url), ''),
                'source_system', 'n8n_chat_histories',
                'source_identity', identity_key
            )),
            created_at,
            created_at,
            coalesce(created_at, now()),
            coalesce(created_at, now()),
            jsonb_build_object('source_table', 'public.n8n_chat_histories', 'source_id', id, 'message', message),
            now()
        from contacts
        on conflict (chatwoot_contact_id) do update set
            name = excluded.name,
            phone_number = coalesce(excluded.phone_number, cw.contacts_current.phone_number),
            email = coalesce(excluded.email, cw.contacts_current.email),
            custom_attributes = cw.contacts_current.custom_attributes || excluded.custom_attributes,
            additional_attributes = cw.contacts_current.additional_attributes || excluded.additional_attributes,
            last_seen_at = greatest(cw.contacts_current.last_seen_at, excluded.last_seen_at),
            raw_payload = excluded.raw_payload,
            updated_at = now();

        with source_rows as (
            select
                h.*,
                nullif(regexp_replace(coalesce(h.celular, ''), '\D', '', 'g'), '') as phone_digits,
                btrim(coalesce(
                    h.message->>'content',
                    h.message->>'text',
                    h.message #>> '{kwargs,content}',
                    h.message #>> '{data,content}',
                    h.message::text,
                    ''
                )) as message_content
            from public.n8n_chat_histories h
            where to_regclass('public.n8n_chat_histories') is not null
        ),
        enriched as (
            select
                *,
                coalesce(
                    'phone:' || nullif(phone_digits, ''),
                    'email:' || nullif(lower(btrim(coalesce(correo, ''))), ''),
                    'name:' || nullif(lower(regexp_replace(btrim(coalesce(nombre, '')), '\s+', '_', 'g')), ''),
                    'session:' || session_id
                ) as identity_key,
                -1000000000::bigint - min(id)::bigint over (partition by session_id) as conversation_id,
                -2000000000::bigint - min(id)::bigint over (
                    partition by coalesce(
                        'phone:' || nullif(phone_digits, ''),
                        'email:' || nullif(lower(btrim(coalesce(correo, ''))), ''),
                        'name:' || nullif(lower(regexp_replace(btrim(coalesce(nombre, '')), '\s+', '_', 'g')), ''),
                        'session:' || session_id
                    )
                ) as contact_id,
                nullif(trim(both '_' from regexp_replace(lower(coalesce(pipeline_etapa, '')), '[^a-z0-9]+', '_', 'g')), '') as stage_label
            from source_rows
        ),
        sessions as (
            select
                session_id,
                min(conversation_id) as conversation_id,
                min(contact_id) as contact_id,
                max(identity_key) as identity_key,
                min(created_at) as first_created_at,
                max(created_at) as last_created_at,
                (array_agg(nombre order by created_at nulls last, id))[1] as nombre,
                (array_agg(correo order by created_at nulls last, id))[1] as correo,
                (array_agg(celular order by created_at nulls last, id))[1] as celular,
                (array_agg(canal order by created_at nulls last, id))[1] as canal,
                (array_agg(perfil_url order by created_at nulls last, id))[1] as perfil_url,
                (array_agg(pipeline_etapa order by created_at desc nulls last, id desc))[1] as pipeline_etapa,
                array_remove(array_agg(distinct stage_label), null) as labels,
                count(*) as total_messages,
                (array_agg(message_content order by created_at desc nulls last, id desc))[1] as last_message_preview,
                jsonb_agg(jsonb_build_object('id', id, 'message', message, 'created_at', created_at) order by created_at, id) as raw_messages
            from enriched
            group by session_id
        )
        insert into cw.conversations_current (
            chatwoot_conversation_id,
            chatwoot_contact_id,
            status,
            labels,
            business_stage_current,
            additional_attributes,
            custom_attributes,
            conversation_custom_attributes,
            contact_custom_attributes,
            meta,
            inbox_name,
            channel_type,
            provider,
            first_message_at,
            last_message_at,
            last_incoming_message_at,
            last_outgoing_message_at,
            created_at_chatwoot,
            updated_at_chatwoot,
            last_activity_at_chatwoot,
            last_non_activity_message_id,
            last_non_activity_message_preview,
            total_messages,
            nombre_completo,
            celular,
            correo,
            canal,
            perfil_url,
            raw_payload,
            source_system,
            external_lead_id,
            import_batch_id,
            imported_at,
            updated_at
        )
        select
            conversation_id,
            contact_id,
            'open',
            coalesce(labels, '{}'::text[]),
            nullif(btrim(pipeline_etapa), ''),
            jsonb_build_object('imported', true, 'source_system', 'n8n_chat_histories', 'session_id', session_id),
            jsonb_strip_nulls(jsonb_build_object(
                'pipeline_etapa', nullif(btrim(pipeline_etapa), ''),
                'nombre_completo', nullif(btrim(nombre), ''),
                'celular', nullif(btrim(celular), ''),
                'correo', nullif(btrim(correo), ''),
                'canal', nullif(btrim(canal), ''),
                'perfil_url', nullif(btrim(perfil_url), ''),
                'source_system', 'n8n_chat_histories',
                'external_lead_id', session_id
            )),
            jsonb_strip_nulls(jsonb_build_object(
                'pipeline_etapa', nullif(btrim(pipeline_etapa), ''),
                'canal', nullif(btrim(canal), ''),
                'perfil_url', nullif(btrim(perfil_url), ''),
                'source_system', 'n8n_chat_histories',
                'external_lead_id', session_id
            )),
            jsonb_strip_nulls(jsonb_build_object(
                'nombre_completo', nullif(btrim(nombre), ''),
                'celular', nullif(btrim(celular), ''),
                'correo', nullif(btrim(correo), ''),
                'canal', nullif(btrim(canal), '')
            )),
            jsonb_build_object(
                'sender', jsonb_strip_nulls(jsonb_build_object(
                    'id', contact_id,
                    'name', coalesce(nullif(btrim(nombre), ''), 'Sin Nombre'),
                    'email', nullif(btrim(correo), ''),
                    'phone_number', nullif(btrim(celular), ''),
                    'identifier', identity_key,
                    'custom_attributes', jsonb_strip_nulls(jsonb_build_object(
                        'nombre_completo', nullif(btrim(nombre), ''),
                        'celular', nullif(btrim(celular), ''),
                        'correo', nullif(btrim(correo), ''),
                        'canal', nullif(btrim(canal), '')
                    ))
                )),
                'assignee', '{}'::jsonb
            ),
            nullif(btrim(canal), ''),
            nullif(btrim(canal), ''),
            'n8n',
            first_created_at,
            last_created_at,
            last_created_at,
            last_created_at,
            first_created_at,
            last_created_at,
            last_created_at,
            null,
            left(coalesce(last_message_preview, ''), 500),
            total_messages,
            nullif(btrim(nombre), ''),
            nullif(btrim(celular), ''),
            nullif(btrim(correo), ''),
            nullif(btrim(canal), ''),
            nullif(btrim(perfil_url), ''),
            jsonb_build_object(
                'source_table', 'public.n8n_chat_histories',
                'session_id', session_id,
                'raw_messages', raw_messages
            ),
            'n8n_chat_histories',
            session_id,
            '00000000-0000-0000-0000-000000000908'::uuid,
            now(),
            now()
        from sessions
        on conflict (chatwoot_conversation_id) do update set
            chatwoot_contact_id = excluded.chatwoot_contact_id,
            labels = excluded.labels,
            business_stage_current = excluded.business_stage_current,
            custom_attributes = cw.conversations_current.custom_attributes || excluded.custom_attributes,
            conversation_custom_attributes = cw.conversations_current.conversation_custom_attributes || excluded.conversation_custom_attributes,
            contact_custom_attributes = cw.conversations_current.contact_custom_attributes || excluded.contact_custom_attributes,
            meta = excluded.meta,
            inbox_name = excluded.inbox_name,
            channel_type = excluded.channel_type,
            last_message_at = excluded.last_message_at,
            last_activity_at_chatwoot = excluded.last_activity_at_chatwoot,
            last_non_activity_message_preview = excluded.last_non_activity_message_preview,
            total_messages = excluded.total_messages,
            raw_payload = excluded.raw_payload,
            updated_at = now();

        with source_rows as (
            select
                h.*,
                nullif(regexp_replace(coalesce(h.celular, ''), '\D', '', 'g'), '') as phone_digits,
                btrim(coalesce(
                    h.message->>'content',
                    h.message->>'text',
                    h.message #>> '{kwargs,content}',
                    h.message #>> '{data,content}',
                    h.message::text,
                    ''
                )) as message_content,
                lower(btrim(coalesce(
                    h.message->>'type',
                    h.message->>'role',
                    h.message #>> '{kwargs,type}',
                    h.message #>> '{data,type}',
                    ''
                ))) as message_kind
            from public.n8n_chat_histories h
            where to_regclass('public.n8n_chat_histories') is not null
        ),
        enriched as (
            select
                *,
                coalesce(
                    'phone:' || nullif(phone_digits, ''),
                    'email:' || nullif(lower(btrim(coalesce(correo, ''))), ''),
                    'name:' || nullif(lower(regexp_replace(btrim(coalesce(nombre, '')), '\s+', '_', 'g')), ''),
                    'session:' || session_id
                ) as identity_key,
                -1000000000::bigint - min(id)::bigint over (partition by session_id) as conversation_id,
                -2000000000::bigint - min(id)::bigint over (
                    partition by coalesce(
                        'phone:' || nullif(phone_digits, ''),
                        'email:' || nullif(lower(btrim(coalesce(correo, ''))), ''),
                        'name:' || nullif(lower(regexp_replace(btrim(coalesce(nombre, '')), '\s+', '_', 'g')), ''),
                        'session:' || session_id
                    )
                ) as contact_id,
                -3000000000::bigint - id::bigint as message_id
            from source_rows
        )
        insert into cw.messages (
            chatwoot_message_id,
            chatwoot_conversation_id,
            chatwoot_contact_id,
            sender_type,
            message_type,
            message_direction,
            content,
            content_type,
            content_attributes,
            additional_attributes,
            sender,
            processed_message_content,
            source_id,
            status,
            created_at_chatwoot,
            updated_at_chatwoot,
            raw_payload
        )
        select
            message_id,
            conversation_id,
            contact_id,
            case
                when message_kind in ('human', 'user', 'incoming') then 'contact'
                when message_kind in ('ai', 'assistant', 'bot', 'outgoing') then 'agent'
                else 'unknown'
            end,
            message_kind,
            case
                when message_kind in ('human', 'user', 'incoming') then 'incoming'
                when message_kind in ('ai', 'assistant', 'bot', 'outgoing') then 'outgoing'
                else 'unknown'
            end,
            message_content,
            'text',
            jsonb_build_object('source_system', 'n8n_chat_histories'),
            jsonb_build_object('session_id', session_id, 'pipeline_etapa', pipeline_etapa),
            jsonb_strip_nulls(jsonb_build_object(
                'id', contact_id,
                'name', nullif(btrim(nombre), ''),
                'email', nullif(btrim(correo), ''),
                'phone_number', nullif(btrim(celular), '')
            )),
            message_content,
            id::text,
            'sent',
            created_at,
            created_at,
            jsonb_build_object('source_table', 'public.n8n_chat_histories', 'source_id', id, 'message', message)
        from enriched
        where message_content <> ''
        on conflict (chatwoot_message_id) do update set
            content = excluded.content,
            message_direction = excluded.message_direction,
            processed_message_content = excluded.processed_message_content,
            raw_payload = excluded.raw_payload;

        insert into cw.raw_ingest (
            source_type,
            endpoint_name,
            event_name,
            entity_type,
            chatwoot_entity_id,
            payload,
            fetched_at,
            processed,
            processing_error
        )
        select
            'manual',
            'n8n_chat_histories',
            'legacy_n8n_history',
            'n8n_chat_history',
            -1000000000::bigint - min(h.id)::bigint over (partition by h.session_id),
            jsonb_build_object(
                'source_table', 'public.n8n_chat_histories',
                'source_id', h.id,
                'session_id', h.session_id,
                'message', h.message,
                'pipeline_etapa', h.pipeline_etapa,
                'nombre', h.nombre,
                'correo', h.correo,
                'celular', h.celular,
                'canal', h.canal,
                'perfil_url', h.perfil_url,
                'created_at', h.created_at
            ),
            coalesce(h.created_at, now()),
            true,
            null
        from public.n8n_chat_histories h
        where to_regclass('public.n8n_chat_histories') is not null
          and not exists (
              select 1
              from cw.raw_ingest existing
              where existing.endpoint_name = 'n8n_chat_histories'
                and existing.payload->>'source_id' = h.id::text
          );

        delete from cw.import_batch_errors
        where import_batch_id = '00000000-0000-0000-0000-000000000908'::uuid;

        insert into cw.import_batch_errors (
            import_batch_id,
            row_number,
            severity,
            field_name,
            reason,
            raw_row
        )
        select
            '00000000-0000-0000-0000-000000000908'::uuid,
            h.id,
            'warning',
            'message',
            'Fila sin contenido textual compatible; se conservó en cw.raw_ingest pero no se creó cw.messages.',
            jsonb_build_object(
                'source_table', 'public.n8n_chat_histories',
                'source_id', h.id,
                'session_id', h.session_id,
                'message', h.message
            )
        from public.n8n_chat_histories h
        where to_regclass('public.n8n_chat_histories') is not null
          and btrim(coalesce(
              h.message->>'content',
              h.message->>'text',
              h.message #>> '{kwargs,content}',
              h.message #>> '{data,content}',
              h.message::text,
              ''
          )) = '';

        do $$
        begin
            if to_regprocedure('cw.refresh_dashboard_discovery(bigint)') is not null then
                perform cw.refresh_dashboard_discovery(0);
            end if;
        end $$;

        -- Validación rápida posterior:
        select 'raw_ingest' as objeto, count(*) from cw.raw_ingest where endpoint_name = 'n8n_chat_histories'
        union all
        select 'contacts_current', count(*) from cw.contacts_current where custom_attributes->>'source_system' = 'n8n_chat_histories'
        union all
        select 'conversations_current', count(*) from cw.conversations_current where source_system = 'n8n_chat_histories'
        union all
        select 'messages', count(*) from cw.messages where content_attributes->>'source_system' = 'n8n_chat_histories'
        union all
        select 'import_batch_errors', count(*) from cw.import_batch_errors where import_batch_id = '00000000-0000-0000-0000-000000000908'::uuid;

        -- ========================================================================
        -- FIN BLOQUE 10: migración opcional public.n8n_chat_histories hacia cw
        -- ========================================================================
        """
    ).strip()


def build_updated_sql() -> str:
    source_sql = read_text(SOURCE_SQL).rstrip()
    source_sql = re.sub(r"-- Fecha:\s*\d{4}-\d{2}-\d{2}\.", f"-- Fecha base: 2026-05-18. Actualización documental: {date.today().isoformat()}.", source_sql)

    blocks = [
        (
            "09: migraciones locales 2026-05-19 y 2026-05-28",
            [
                "20260519130831_auto_discover_chatwoot_catalogs.sql",
                "20260519133145_centralize_dashboard_settings_in_cw.sql",
                "20260528145602_meta_ads_insights_cache.sql",
                "20260528152534_meta_ads_request_range_columns.sql",
            ],
        ),
    ]

    appended: list[str] = [
        "",
        "-- ========================================================================",
        "-- ADDENDUM 2026-05-28: estado vigente del repo",
        "-- Incluye catálogos automáticos, centralización de settings en cw, Meta Ads",
        "-- y migración opcional desde public.n8n_chat_histories.",
        "-- ========================================================================",
    ]

    for title, migrations in blocks:
        appended.append("")
        appended.append(f"-- ========================================================================")
        appended.append(f"-- INICIO BLOQUE {title}")
        appended.append("-- ========================================================================")
        for migration in migrations:
            appended.append("")
            appended.append(f"-- =============================================================")
            appended.append(f"-- MIGRACION LOCAL: {migration}")
            appended.append(f"-- =============================================================")
            appended.append(migration_text(migration))
        appended.append("")
        appended.append(f"-- ========================================================================")
        appended.append(f"-- FIN BLOQUE {title}")
        appended.append(f"-- ========================================================================")

    appended.append("")
    appended.append(n8n_chat_histories_migration_sql())
    appended.append("")
    appended.append(
        dedent(
            """
            -- ========================================================================
            -- INICIO BLOQUE 11: validaciones finales addendum 2026-05-28
            -- ========================================================================

            select to_regclass('public.dashboard_tag_settings') as public_dashboard_tag_settings,
                   to_regclass('cw.dashboard_tag_settings') as cw_dashboard_tag_settings,
                   to_regclass('cw.label_catalog') as label_catalog,
                   to_regclass('cw.attribute_key_catalog') as attribute_key_catalog,
                   to_regclass('cw.meta_campaigns_current') as meta_campaigns_current,
                   to_regclass('cw.meta_adset_insights_cache') as meta_adset_insights_cache,
                   to_regclass('cw.meta_ads_sync_runs') as meta_ads_sync_runs;

            select table_schema, table_name
            from information_schema.tables
            where table_schema = 'cw'
              and table_name in (
                  'dashboard_tag_settings',
                  'label_catalog',
                  'attribute_key_catalog',
                  'meta_campaigns_current',
                  'meta_adset_insights_cache',
                  'meta_ads_sync_runs',
                  'import_batches',
                  'import_batch_errors'
              )
            order by table_name;

            select schemaname, tablename, rowsecurity
            from pg_tables
            where schemaname = 'cw'
              and tablename in (
                  'label_catalog',
                  'attribute_key_catalog',
                  'meta_campaigns_current',
                  'meta_adset_insights_cache',
                  'meta_ads_sync_runs'
              )
            order by tablename;

            -- ========================================================================
            -- FIN BLOQUE 11: validaciones finales addendum 2026-05-28
            -- ========================================================================
            """
        ).strip()
    )

    return source_sql + "\n" + "\n".join(appended).rstrip() + "\n"


def generate_updated_sql() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    UPDATED_SQL.write_text(build_updated_sql(), encoding="utf-8")


def build_technical_doc() -> None:
    if SOURCE_TECH_DOCX.exists():
        document = Document(SOURCE_TECH_DOCX)
        document.add_section(WD_SECTION.NEW_PAGE)
    else:
        document = Document()
        configure_document(
            document,
            "IMP-SUPA-SIMPLIA-003 | v2.1 | Documento controlado",
            "Implementación Supabase SimpliaLeads ISO 10013",
        )
        add_title(
            document,
            "Implementación Supabase SimpliaLeads",
            "Runbook técnico actualizado",
        )

    document.add_heading("Actualización v2.1 - Estado vigente al 28/05/2026", level=1)
    add_callout(
        document,
        "Conclusión de auditoría documental",
        "El Word y SQL base fechados el 18/05/2026 no cubren todo lo que hoy usa el repo. Esta actualización agrega el estado posterior: migraciones del 19/05 para discovery/settings y migraciones del 28/05 para Meta Ads, además de la ruta opcional para migrar public.n8n_chat_histories hacia cw.",
        "yellow",
    )

    add_table(
        document,
        [
            ["Área", "Documento 18/05/2026", "Estado vigente del repo", "Acción v2.1"],
            ["Settings dashboard", "Incluye public.dashboard_tag_settings como fallback.", "El código lee cw.dashboard_tag_settings. public.dashboard_tag_settings se elimina con 20260519133145.", "Marcar cw como única fuente y validar que public quede null."],
            ["Catálogos Chatwoot", "No incluye label_catalog ni attribute_key_catalog.", "El repo usa catálogos automáticos para etiquetas y atributos detectados.", "Agregar 20260519130831 y función cw.refresh_dashboard_discovery."],
            ["Meta Ads", "No contemplado.", "TrendLayer y exportes usan meta-campaign-insights y caché en cw.", "Agregar función, tablas, secrets y pruebas."],
            ["OpenAI", "Incluye OPENAI_API_KEY básico.", "El repo soporta overrides por perfil/formato, timeouts y max tokens.", "Documentar variables adicionales server-side."],
            ["Histórico n8n", "Excluía tablas public de n8n.", "El onboarding requiere migrar contenido útil de public.n8n_chat_histories.", "Agregar bloque opcional hacia raw_ingest, contacts, conversations, messages y batch/errors."],
        ],
        [1.2, 1.8, 2.0, 2.0],
    )

    document.add_heading("Objetos agregados por esta versión", level=2)
    add_table(
        document,
        [
            ["Objeto", "Tipo", "Uso operativo"],
            ["cw.label_catalog", "Tabla", "Catálogo automático de etiquetas reales de Chatwoot vistas por sync/webhook."],
            ["cw.attribute_key_catalog", "Tabla", "Catálogo automático de llaves de atributos vistas en definiciones y payloads."],
            ["cw.refresh_dashboard_discovery(bigint)", "Función", "Actualiza catálogos y sugerencias sin decidir qué cuenta como venta/cita/SQL."],
            ["cw.meta_campaigns_current", "Tabla", "Snapshot de campañas de Meta Ads por ad account."],
            ["cw.meta_adset_insights_cache", "Tabla", "Caché de insights por rango solicitado y conjunto de anuncios."],
            ["cw.meta_ads_sync_runs", "Tabla", "Bitácora de requests Meta Ads, cache hits, errores y rate-limit metadata."],
            ["meta-campaign-insights", "Edge Function", "Consulta Meta Graph API, escribe caché en cw y devuelve filas al dashboard."],
        ],
        [2.1, 1.2, 4.0],
    )

    document.add_heading("Snapshot de Producción verificado (29/05/2026)", level=2)
    add_table(
        document,
        [
            ["Componente", "Estado verificado", "Fuente de verificación"],
            [
                "Edge Functions",
                "chatwoot-sync, chatwoot-repair-conversations, chatwoot-label-webhook, generate-ai-report, send-scheduled-reports, meta-campaign-insights (ACTIVE).",
                "Listado real de funciones en proyecto knsmqbkdsfhttizaepzv.",
            ],
            [
                "Secrets Chatwoot",
                "CHATWOOT_BASE_URL, CHATWOOT_ACCOUNT_ID, CHATWOOT_API_TOKEN, CHATWOOT_WEBHOOK_SECRET, VITE_CHATWOOT_BASE_URL, VITE_CHATWOOT_ACCOUNT_ID, VITE_CHATWOOT_API_TOKEN.",
                "Pantalla de Edge Functions > Secrets compartida por el usuario.",
            ],
            [
                "Secrets IA/Reportes",
                "OPENAI_API_KEY, OPENAI_REPORT_MODEL, OPENAI_REPORT_REASONING_EFFORT, RESEND_API_KEY, RESEND_FROM_EMAIL.",
                "Pantalla de Edge Functions > Secrets compartida por el usuario.",
            ],
            [
                "Secrets Supabase runtime",
                "SUPABASE_URL, SUPABASE_ANON_KEY, SUPABASE_SERVICE_ROLE_KEY, SUPABASE_DB_URL, SUPABASE_JWKS, SUPABASE_PUBLISHABLE_KEYS, SUPABASE_SECRET_KEYS.",
                "CLI `supabase secrets list` ejecutado por el usuario (29/05/2026).",
            ],
            [
                "Secrets Meta Ads",
                "META_AD_ACCOUNT_ID, META_SYSTEM_USER_TOKEN, META_GRAPH_API_VERSION, META_CACHE_TTL_SECONDS.",
                "CLI `supabase secrets list` + pantalla de Edge Functions (actualizados 28/05/2026).",
            ],
        ],
        [1.8, 2.8, 2.7],
    )
    add_callout(
        document,
        "Nota de trazabilidad",
        "La lectura directa de secrets por CLI se validó el 29/05/2026. El `db pull` remoto no se pudo completar en ese momento porque faltaba Docker Desktop en la máquina que ejecutó el comando, así que el estado de esquema se dejó respaldado con migraciones del repo y validaciones SQL incluidas en el documento.",
        "yellow",
    )

    document.add_heading("Secretos vigentes", level=2)
    add_table(
        document,
        [
            ["Secret", "Obligatorio", "Dónde vive", "Uso"],
            ["OPENAI_API_KEY", "Sí para reportes IA", "Supabase Edge Function Secrets", "generate-ai-report y send-scheduled-reports."],
            ["OPENAI_REPORT_MODEL", "Opcional", "Supabase Edge Function Secrets", "Default compatible con repo: gpt-5.4-mini."],
            ["OPENAI_REPORT_MODEL_PDF", "Opcional", "Supabase Edge Function Secrets", "Override para reportes PDF."],
            ["OPENAI_REPORT_MODEL_TABLE", "Opcional", "Supabase Edge Function Secrets", "Override para reportes tabulares."],
            ["OPENAI_REPORT_MODEL_DAILY_OPERATIONS_EXCEL", "Opcional", "Supabase Edge Function Secrets", "Override por perfil/formato."],
            ["OPENAI_REPORT_TIMEOUT_MS", "Opcional", "Supabase Edge Function Secrets", "Timeout de generación síncrona."],
            ["OPENAI_REPORT_START_TIMEOUT_MS", "Opcional", "Supabase Edge Function Secrets", "Timeout de inicio para background."],
            ["OPENAI_REPORT_MAX_OUTPUT_TOKENS", "Opcional", "Supabase Edge Function Secrets", "Límite de salida de reportes IA."],
            ["META_AD_ACCOUNT_ID", "Sí si se activa Meta Ads", "Supabase Edge Function Secrets", "Ad account sin prefijo act_."],
            ["META_SYSTEM_USER_TOKEN", "Sí si se activa Meta Ads", "Supabase Edge Function Secrets", "Token de system user para Graph API."],
            ["META_GRAPH_API_VERSION", "Opcional", "Supabase Edge Function Secrets", "Default repo: v20.0."],
            ["META_CACHE_TTL_SECONDS", "Opcional", "Supabase Edge Function Secrets", "Default repo: 900 segundos."],
        ],
        [2.1, 1.1, 2.0, 2.4],
    )

    document.add_heading("SQL consolidado actualizado", level=2)
    add_bullets(
        document,
        [
            f"El SQL actualizado queda en {UPDATED_SQL.relative_to(REPO_ROOT)}.",
            "El archivo conserva el SQL base y agrega un addendum con migraciones 20260519130831, 20260519133145, 20260528145602 y 20260528152534.",
            "El bloque opcional de n8n no elimina public.n8n_chat_histories; solo la lee, migra lo compatible y deja errores en cw.import_batch_errors.",
            "Ejecutar el bloque n8n únicamente en staging primero y comparar conteos antes de repetir en producción.",
        ],
    )

    document.add_heading("Migración opcional public.n8n_chat_histories", level=2)
    add_table(
        document,
        [
            ["Origen", "Destino cw", "Regla de mapeo"],
            ["id, session_id, message, created_at", "cw.raw_ingest", "Se conserva cada fila como payload original auditado."],
            ["celular/correo/nombre/session_id", "cw.contacts_current", "Identidad en orden: celular, correo, nombre, session_id."],
            ["session_id", "cw.conversations_current", "Una conversación por sesión, con IDs negativos determinísticos para no chocar con Chatwoot."],
            ["message jsonb", "cw.messages", "Contenido desde message.content, message.text, message.kwargs.content, message.data.content o JSON completo."],
            ["pipeline_etapa", "labels/business_stage_current/custom_attributes", "Se normaliza como etiqueta y se conserva el texto original."],
            ["filas sin contenido textual", "cw.import_batch_errors", "Quedan como warning y no generan cw.messages."],
        ],
        [1.8, 2.0, 3.4],
    )

    document.add_heading("Validaciones obligatorias", level=2)
    add_table(
        document,
        [
            ["Validación", "Consulta / acción", "Éxito esperado"],
            ["Settings final", "select to_regclass('public.dashboard_tag_settings'), to_regclass('cw.dashboard_tag_settings')", "public null; cw existe."],
            ["Catálogos", "select count(*) from cw.label_catalog; select count(*) from cw.attribute_key_catalog;", "Conteos > 0 después del sync si Chatwoot tiene datos."],
            ["Meta Ads", "Invocar meta-campaign-insights con rango válido.", "Respuesta ok=true o error controlado por secrets faltantes."],
            ["n8n staging", "Comparar count(*) public.n8n_chat_histories vs cw.raw_ingest endpoint n8n_chat_histories.", "cw.raw_ingest conserva todas las filas procesadas."],
            ["RLS", "select schemaname, tablename, rowsecurity from pg_tables where schemaname='cw';", "Tablas sensibles con rowsecurity=true."],
        ],
        [1.6, 3.2, 2.2],
    )

    document.add_heading("Fuentes oficiales usadas para esta actualización", level=2)
    add_table(
        document,
        [
            ["Fuente", "URL", "Uso"],
            ["Chatwoot pricing", "https://www.chatwoot.com/pricing/", "Plan Startups como mínimo viable y criterios de upgrade."],
            ["Chatwoot WhatsApp Cloud", "https://www.chatwoot.com/docs/product/channels/whatsapp/whatsapp-cloud", "Embedded Signup recomendado y manual setup como alternativa."],
            ["Chatwoot WhatsApp Embedded Signup", "https://www.chatwoot.com/hc/user-guide/articles/1752129193-how-to-use-whatsapp-embedded-signup", "Prerrequisitos y flujo dentro de Chatwoot Cloud."],
            ["Chatwoot Instagram", "https://www.chatwoot.com/hc/user-guide/articles/1744361165-how-to-setup-an-instagram-channel-via-instagram-login", "Conexión Instagram Business Login."],
            ["Chatwoot TikTok", "https://www.chatwoot.com/hc/user-guide/en/categories/other-channels", "Limitaciones del canal TikTok Business Messaging."],
            ["Supabase seguridad API", "https://supabase.com/docs/guides/api/securing-your-api", "RLS y schemas expuestos por Data API."],
            ["OpenAI models", "https://developers.openai.com/api/docs/models/all", "Modelo gpt-5.4-mini disponible para API."],
        ],
        [1.7, 2.7, 2.7],
    )

    document.core_properties.title = "Implementación Supabase SimpliaLeads ISO 10013 v2.1"
    document.core_properties.author = "Simplia"
    document.core_properties.subject = "Actualización documental, SQL consolidado, Meta Ads y migración n8n"
    document.core_properties.keywords = "Supabase, Chatwoot, ISO 10013, SimpliaLeads, Meta Ads, OpenAI, n8n"
    document.save(UPDATED_TECH_DOCX)


def build_onboarding_doc() -> None:
    document = Document()
    configure_document(
        document,
        "ONB-CW-CLOUD-001 | v1.0 | ISO 10013",
        "Onboarding Migración Chatwoot Cloud - Simplia",
    )
    add_title(
        document,
        "Onboarding de Migración a Chatwoot Cloud",
        "Procedimiento ISO 10013 con controles operativos ISO 9001 para canales, Supabase, OpenAI, Resend y Meta Ads",
    )

    document.add_heading("1. Control documental", level=1)
    add_table(
        document,
        [
            ["Campo", "Valor"],
            ["Código", "ONB-CW-CLOUD-001"],
            ["Versión", "1.0"],
            ["Fecha", DOC_DATE],
            ["Estado", "Borrador controlado para ejecución"],
            ["Modelo documental", "ISO 10013 / información documentada. Incluye controles de proceso ISO 9001."],
            ["Alcance", "Migrar operación desde Chatwoot Self-hosted hacia Chatwoot Cloud y conectar SimpliaLeads/Supabase."],
            ["Exclusión", "No migra la base interna de Chatwoot Self-hosted. El histórico útil se normaliza desde public.n8n_chat_histories hacia cw."],
        ],
        [2.1, 5.0],
    )

    document.add_heading("2. Resumen ejecutivo", level=1)
    add_callout(
        document,
        "Decisión operativa",
        "El proyecto deja de depender de Chatwoot Self-hosted para la atención diaria. Chatwoot Cloud queda como plataforma operativa, mientras Supabase conserva el histórico útil y alimenta el dashboard.",
        "green",
    )
    add_bullets(
        document,
        [
            "Plan Chatwoot recomendado: Startups como mínimo viable; validar Business si el cliente necesita Teams, Automation Rules o mayor retención.",
            "Canales a reconectar en Cloud: WhatsApp, Facebook/Messenger, Instagram y TikTok cuando esté disponible para la cuenta.",
            "Supabase se replica con el SQL actualizado y las Edge Functions del repo.",
            "OpenAI y Resend se configuran como secrets server-side.",
            "Meta Ads se activa solo si el cliente entrega ad account y token de system user.",
        ],
    )

    document.add_heading("3. Responsables y evidencias", level=1)
    add_table(
        document,
        [
            ["Rol", "Responsabilidad", "Evidencia requerida"],
            ["Cliente / dueño del negocio", "Aprobar plan Chatwoot, pagos, permisos Meta y acceso a canales.", "Confirmación escrita, método de pago, capturas de activos."],
            ["Administrador Meta", "Autorizar WhatsApp, Facebook, Instagram y TikTok.", "Capturas de Business Manager, WABA, páginas, IG Business y permisos."],
            ["Técnico Simplia", "Ejecutar SQL, secrets, deploy de funciones y pruebas.", "Logs, capturas Supabase, conteos SQL y pruebas funcionales."],
            ["Administrador Chatwoot", "Crear workspace, agentes, inboxes y tokens API.", "Captura workspace, plan activo, inboxes y API token guardado."],
            ["Responsable de datos", "Validar migración n8n y conteos de staging.", "Acta de validación, conteos origen/destino y errores revisados."],
        ],
        [1.4, 3.2, 2.6],
    )

    document.add_heading("4. Entradas necesarias antes de iniciar", level=1)
    add_table(
        document,
        [
            ["Entrada", "Obligatorio", "Notas"],
            ["Correo empresarial para Chatwoot Cloud", "Sí", "Debe poder recibir invitaciones y facturación."],
            ["Método de pago Chatwoot", "Sí", "Contratar Startups salvo decisión distinta."],
            ["Acceso Meta Business", "Sí para canales Meta", "Debe administrar WABA, página Facebook e Instagram Business."],
            ["TikTok Business Account", "Solo si aplica", "Debe estar en región elegible y aceptar mensajes directos de todos."],
            ["Supabase PROJECT_REF, anon key y service role", "Sí", "Service role nunca se pega en frontend."],
            ["Chatwoot Cloud Account ID y API token", "Sí", "Token admin/owner para sync y repair."],
            ["OPENAI_API_KEY", "Sí para reportes IA", "Secret server-side; nunca VITE_*."],
            ["RESEND_API_KEY y remitente verificado", "Sí para reportes por correo", "Dominio o remitente debe estar aprobado en Resend."],
            ["META_AD_ACCOUNT_ID y META_SYSTEM_USER_TOKEN", "Solo si aplica", "Necesario para Meta Ads Insights."],
        ],
        [2.3, 1.1, 3.6],
    )

    document.add_heading("5. Procedimiento de migración", level=1)
    add_numbered(
        document,
        [
            "Congelar ventana de cambio: definir fecha/hora, responsable y canales a migrar.",
            "Respaldar evidencia de Self-hosted: URL, canales conectados, webhooks, agentes, etiquetas, atributos y automatizaciones n8n.",
            "Crear workspace Chatwoot Cloud y contratar plan Startups; Business solo si los criterios de upgrade lo justifican.",
            "Crear agentes, permisos e inboxes base en Chatwoot Cloud.",
            "Desconectar o pausar canales en Self-hosted para evitar doble recepción.",
            "Reconectar WhatsApp en Cloud con Embedded Signup cuando aplique; usar manual setup solo si el número o el proveedor lo exige.",
            "Reconectar Facebook/Messenger autorizando la página correcta.",
            "Reconectar Instagram con Instagram Business Login y validar permisos de mensajes.",
            "Validar TikTok Business Messaging; si se activa, conectar desde Settings > Inboxes > TikTok y documentar limitaciones.",
            "Generar API token en Chatwoot Cloud y configurar Supabase Edge Function Secrets.",
            "Ejecutar el SQL consolidado actualizado en Supabase y desplegar Edge Functions.",
            "Ejecutar migración opcional de public.n8n_chat_histories solo en staging; aprobar conteos antes de producción.",
            "Configurar webhook de Chatwoot hacia chatwoot-label-webhook con CHATWOOT_WEBHOOK_SECRET.",
            "Ejecutar sync inicial, repair si aplica, reportes IA y reportes programados.",
            "Cerrar migración con evidencia, riesgos residuales y aprobación del cliente.",
        ],
    )

    document.add_heading("6. Canales Chatwoot Cloud", level=1)
    add_table(
        document,
        [
            ["Canal", "Ruta recomendada", "Prueba mínima", "Notas"],
            ["WhatsApp", "Settings > Inboxes > Add Inbox > WhatsApp > Embedded Signup.", "Mensaje entrante y respuesta desde Chatwoot.", "Embedded Signup auto-configura webhooks/tokens si aplica."],
            ["Facebook/Messenger", "Settings > Inboxes > Add Inbox > Messenger.", "Mensaje desde página a inbox.", "Autorizar con cuenta admin de la página."],
            ["Instagram", "Settings > Inboxes > Add Inbox > Instagram.", "DM entrante y respuesta.", "Cuenta debe ser Instagram Business."],
            ["TikTok", "Settings > Inboxes > Add Inbox > TikTok.", "DM entrante y respuesta si está habilitado.", "Limitaciones: cliente inicia conversación, 10 mensajes antes de respuesta, ventana 48h, soporta texto/imagen/post share."],
        ],
        [1.3, 2.6, 1.8, 2.4],
    )

    document.add_heading("7. Supabase y datos", level=1)
    add_bullets(
        document,
        [
            f"Usar el SQL actualizado: {UPDATED_SQL.relative_to(REPO_ROOT)}.",
            "Aplicar primero bloques base, luego addendum 2026-05-28.",
            "La fuente vigente de settings es cw.dashboard_tag_settings; public.dashboard_tag_settings no debe quedar operativa.",
            "El bloque public.n8n_chat_histories no borra la tabla origen; migra lo compatible y deja warnings en cw.import_batch_errors.",
            "Después de sync o migración n8n, ejecutar cw.refresh_dashboard_discovery(0) si existe.",
        ],
    )
    add_table(
        document,
        [
            ["Dato histórico", "Destino", "Criterio de aceptación"],
            ["Filas public.n8n_chat_histories", "cw.raw_ingest", "Conteo destino igual a filas procesadas del origen."],
            ["Sesiones", "cw.conversations_current", "Una conversación por session_id."],
            ["Contactos", "cw.contacts_current", "Dedupe por celular/correo/nombre/session_id."],
            ["Mensajes compatibles", "cw.messages", "Contenido visible desde message.content/text/kwargs/data o JSON."],
            ["Incompatibles", "cw.import_batch_errors", "Warnings revisados, sin bloquear conservación en raw_ingest."],
        ],
        [1.8, 2.2, 3.2],
    )

    document.add_heading("8. Secrets y despliegue", level=1)
    add_code(
        document,
        r"""
        npx supabase secrets set `
          VITE_CHATWOOT_BASE_URL="https://app.chatwoot.com" `
          VITE_CHATWOOT_ACCOUNT_ID="<CHATWOOT_ACCOUNT_ID>" `
          VITE_CHATWOOT_API_TOKEN="<CHATWOOT_API_TOKEN>" `
          CHATWOOT_BASE_URL="https://app.chatwoot.com" `
          CHATWOOT_ACCOUNT_ID="<CHATWOOT_ACCOUNT_ID>" `
          CHATWOOT_API_TOKEN="<CHATWOOT_API_TOKEN>" `
          CHATWOOT_WEBHOOK_SECRET="<CHATWOOT_WEBHOOK_SECRET>" `
          OPENAI_API_KEY="<OPENAI_API_KEY>" `
          OPENAI_REPORT_MODEL="gpt-5.4-mini" `
          OPENAI_REPORT_REASONING_EFFORT="low" `
          RESEND_API_KEY="<RESEND_API_KEY>" `
          RESEND_FROM_EMAIL="Simplia Leads <reportes@tu-dominio.com>" `
          META_AD_ACCOUNT_ID="<META_AD_ACCOUNT_ID>" `
          META_SYSTEM_USER_TOKEN="<META_SYSTEM_USER_TOKEN>" `
          META_GRAPH_API_VERSION="v20.0" `
          META_CACHE_TTL_SECONDS="900" `
          --project-ref "<PROJECT_REF>"

        npx supabase functions deploy chatwoot-sync chatwoot-repair-conversations generate-ai-report send-scheduled-reports meta-campaign-insights --project-ref "<PROJECT_REF>" --use-api
        npx supabase functions deploy chatwoot-label-webhook --project-ref "<PROJECT_REF>" --use-api --no-verify-jwt
        """,
        "Comandos base",
    )

    document.add_heading("9. Validación final", level=1)
    add_table(
        document,
        [
            ["Prueba", "Resultado esperado", "Evidencia"],
            ["Login dashboard", "Usuarios entran con rol correcto.", "Captura sin exponer tokens."],
            ["Sync Chatwoot", "cw.sync_runs success y conversaciones visibles.", "SQL count + log función."],
            ["Webhook Chatwoot", "Cambio de etiqueta genera evento o actualiza snapshot.", "Log 200 + fila cw."],
            ["Reporte IA", "generate-ai-report responde sin Missing OPENAI_API_KEY.", "Captura reporte o response ok."],
            ["Reporte programado", "send-scheduled-reports registra run y Resend acepta.", "cw.automated_report_runs + email de prueba."],
            ["Meta Ads", "meta-campaign-insights ok=true si secrets fueron entregados.", "Response o log de error controlado."],
            ["n8n histórico", "Conteos origen/destino revisados.", "Tabla de conteos firmada."],
        ],
        [1.5, 3.3, 2.2],
    )

    document.add_heading("10. Riesgos y mitigaciones", level=1)
    add_table(
        document,
        [
            ["Riesgo", "Impacto", "Mitigación"],
            ["Doble recepción durante cambio", "Mensajes duplicados o perdidos.", "Congelar ventana, pausar Self-hosted y probar Cloud antes de abrir tráfico."],
            ["Permisos Meta insuficientes", "No se conectan canales.", "Validar admin real del Business/activos antes del día de cambio."],
            ["Token Chatwoot sin permisos", "Sync incompleto.", "Generar token con usuario admin/owner y probar /inboxes, /labels, /custom_attribute_definitions."],
            ["n8n histórico no interpretable", "Mensajes incompletos en cw.messages.", "Conservar todo en cw.raw_ingest y revisar warnings."],
            ["Secrets expuestos", "Riesgo crítico de seguridad.", "No usar VITE_* para OpenAI/Resend/Meta/service role; rotar si se expone."],
            ["Meta Ads sin token válido", "Módulo Meta no carga.", "Marcar N.A. hasta recibir ad account/token."],
        ],
        [1.8, 2.2, 3.2],
    )

    document.add_heading("11. Fuentes oficiales", level=1)
    add_table(
        document,
        [
            ["Fuente", "URL", "Uso"],
            ["Chatwoot pricing", "https://www.chatwoot.com/pricing/", "Plan Startups y criterios de upgrade."],
            ["Crear cuenta Chatwoot", "https://www.chatwoot.com/hc/user-guide/articles/1677242820-create-your-chatwoot-account", "Alta de workspace Cloud."],
            ["WhatsApp Cloud", "https://www.chatwoot.com/docs/product/channels/whatsapp/whatsapp-cloud", "Embedded Signup vs manual setup."],
            ["WhatsApp Embedded Signup", "https://www.chatwoot.com/hc/user-guide/articles/1752129193-how-to-use-whatsapp-embedded-signup", "Flujo recomendado en Cloud."],
            ["Instagram", "https://www.chatwoot.com/hc/user-guide/articles/1744361165-how-to-setup-an-instagram-channel-via-instagram-login", "Instagram Business Login."],
            ["TikTok", "https://www.chatwoot.com/hc/user-guide/en/categories/other-channels", "Conexión y limitaciones de TikTok."],
            ["Supabase securing API", "https://supabase.com/docs/guides/api/securing-your-api", "RLS y Data API."],
            ["OpenAI models", "https://developers.openai.com/api/docs/models/all", "Modelo gpt-5.4-mini."],
        ],
        [1.5, 2.9, 2.8],
    )

    document.add_heading("12. Aprobación", level=1)
    add_table(
        document,
        [
            ["Elaborado por", "Revisado por", "Aprobado por"],
            ["Simplia / Técnico", "", ""],
            ["Firma / fecha", "Firma / fecha", "Firma / fecha"],
        ],
        [2.2, 2.2, 2.2],
    )

    document.core_properties.title = "Onboarding Migración Chatwoot Cloud ISO 10013"
    document.core_properties.author = "Simplia"
    document.core_properties.subject = "Migración desde Chatwoot Self-hosted a Chatwoot Cloud"
    document.core_properties.keywords = "Chatwoot Cloud, ISO 10013, Supabase, n8n, OpenAI, Meta Ads"
    document.save(ONBOARDING_DOCX)


def main() -> None:
    if not SOURCE_SQL.exists():
        raise FileNotFoundError(SOURCE_SQL)
    generate_updated_sql()
    build_technical_doc()
    build_onboarding_doc()
    print(f"SQL={UPDATED_SQL}")
    print(f"TECH_DOCX={UPDATED_TECH_DOCX}")
    print(f"ONBOARDING_DOCX={ONBOARDING_DOCX}")


if __name__ == "__main__":
    main()
